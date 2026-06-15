# -*- coding: utf-8 -*-
"""Markdown -> Word(.docx) 変換(kiduri 資料用)。
見出し/段落/太字/コード/表/箇条書き/番号/引用/水平線/図プレースホルダ/簡易数式整形に対応。
figures/<id>.png があれば図として埋め込み、無ければプレースホルダ枠を出す。
"""
import os, re, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG_DIR = os.path.join(HERE, "figures")
JP_FONT = "Yu Gothic"
MONO_FONT = "Consolas"

def set_run_font(run, name=JP_FONT, size=10.5, bold=False, italic=False, color=None, mono=False):
    fn = MONO_FONT if mono else name
    run.font.name = fn
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:eastAsia'), JP_FONT if not mono else MONO_FONT)
    rfonts.set(qn('w:ascii'), fn)
    rfonts.set(qn('w:hAnsi'), fn)

def shade(paragraph, fill="F2F3F5"):
    pPr = paragraph._p.get_or_add_pPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), fill)
    pPr.append(sh)

def box_border(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        e = OxmlElement('w:'+side)
        e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'6'); e.set(qn('w:space'),'6'); e.set(qn('w:color'),'B8BCC4')
        pbdr.append(e)
    pPr.append(pbdr)

GREEK = {r'\\Theta':'Θ',r'\\theta':'θ',r'\\tau':'τ',r'\\delta':'δ',r'\\Delta':'Δ',
         r'\\Gamma':'Γ',r'\\gamma':'γ',r'\\kappa':'κ',r'\\sigma':'σ',r'\\Sigma':'Σ',
         r'\\Pi':'Π',r'\\pi':'π',r'\\phi':'φ',r'\\varphi':'φ',r'\\mu':'μ',r'\\lambda':'λ',
         r'\\alpha':'α',r'\\beta':'β',r'\\epsilon':'ε'}
SYM = {r'\\Vert':'‖',r'\\mid':'|',r'\\in':'∈',r'\\ge':'≥',r'\\le':'≥'.replace('≥','≤'),
       r'\\times':'×',r'\\cdot':'·',r'\\approx':'≈',r'\\to':'→',r'\\oplus':'⊕',
       r'\\varnothing':'∅',r'\\max':'max',r'\\min':'min',r'\\sim':'~',r'\\circ':'∘',
       r'\\langle':'⟨',r'\\rangle':'⟩',r'\\sum':'Σ',r'\\#':'#',
       r'\\,':' ',r'\\;':' ',r'\\!':'',r'\\ ':' '}

def clean_math(t):
    # 区切り
    t = t.replace('\\[','').replace('\\]','').replace('\\(','').replace('\\)','')
    # \text/\mathrm 等を先に展開(\frac の入れ子対策。複数回)
    for _ in range(3):
        t = re.sub(r'\\(text|mathrm|mathbf|mathit|mathbb|operatorname)\{([^{}]*)\}', r'\2', t)
    # \frac(入れ子解消後に)
    for _ in range(2):
        t = re.sub(r'\\frac\{([^{}]*)\}\{([^{}]*)\}', r'(\1)/(\2)', t)
    t = re.sub(r'\\overline\{([^{}]*)\}', lambda m: m.group(1)+'̄', t)
    # 大型作用素は先に記号化(後段の \big サイズ除去が \bigoplus 等を壊さないように)
    for a,b in [('\\bigoplus','⊕'),('\\bigotimes','⊗'),('\\bigcup','∪'),('\\bigcap','∩'),
                ('\\bigvee','∨'),('\\bigwedge','∧'),('\\bigsqcup','⊔')]:
        t = t.replace(a,b)
    # サイズ・区切り指定は除去(\big( → ( )
    for cmd in ['biggl','biggr','bigg','Biggl','Biggr','Bigg','bigl','bigr','Bigl','Bigr','big','Big','left','right']:
        t = t.replace('\\'+cmd, '')
    # 記号(語を生むものを含む)を先にマップ
    for a,b in [('\\bigoplus','⊕'),('\\oplus','⊕'),('\\Vert','‖'),('\\|','‖'),('\\langle','⟨'),('\\rangle','⟩'),
                ('\\varnothing','∅'),('\\approx','≈'),('\\times','×'),('\\cdot','·'),('\\circ','∘'),
                ('\\mid','|'),('\\sum','Σ'),('\\max','max'),('\\min','min'),('\\sim','〜'),
                ('\\ge','≥'),('\\le','≤'),('\\in','∈'),('\\to','→'),('\\#','#'),
                ('\\,',' '),('\\;',' '),('\\:',' '),('\\!',''),('\\>',' '),('\\ ',' ')]:
        t = t.replace(a,b)
    # ギリシャ
    for k,v in [('varphi','φ'),('Theta','Θ'),('theta','θ'),('tau','τ'),('Delta','Δ'),('delta','δ'),
                ('Gamma','Γ'),('gamma','γ'),('kappa','κ'),('Sigma','Σ'),('sigma','σ'),('Pi','Π'),
                ('pi','π'),('phi','φ'),('mu','μ'),('lambda','λ'),('alpha','α'),('beta','β'),('epsilon','ε')]:
        t = t.replace('\\'+k, v)
    # エスケープ文字
    t = t.replace('\\_','_').replace('\\&','&').replace('\\%','%').replace('\\$','$')
    t = t.replace('\\{','{').replace('\\}','}').replace('\\\\',' ')
    # 下付き・上付きの波括弧を外す(u_{dom}→u_dom, ^{-1}→^-1)
    t = re.sub(r'_\{([^{}]*)\}', r'_\1', t)
    t = re.sub(r'\^\{([^{}]*)\}', r'^\1', t)
    # 残った未知コマンドは名前を残さず除去
    t = re.sub(r'\\[A-Za-z]+','',t)
    t = re.sub(r'[ \t]{2,}',' ',t)
    return t.strip()

def add_inline(paragraph, text, do_math=False, base_size=10.5):
    if do_math:
        text = clean_math(text)
    # トークン分解: `code`, **bold**, [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)  # リンクはテキストのみ
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()]); set_run_font(r, size=base_size)
        tok = m.group(0)
        if tok.startswith('`'):
            r = paragraph.add_run(tok[1:-1]); set_run_font(r, size=base_size-1, mono=True, color=(0xB0,0x30,0x30))
        else:
            r = paragraph.add_run(tok[2:-2]); set_run_font(r, size=base_size, bold=True)
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:]); set_run_font(r, size=base_size)

def add_figure(doc, fid, caption):
    path = os.path.join(FIG_DIR, fid + ".png")
    if not os.path.exists(path):
        path = os.path.join(FIG_DIR, fid + ".jpg")
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(6.3))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cap, caption, base_size=9)
    else:
        p = doc.add_paragraph(); box_border(p); shade(p, "F7F8FA"); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("［スクリーンショット挿入位置］"); set_run_font(r, size=10, bold=True, color=(0x88,0x8E,0x99))
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_inline(cap, caption + "（個人情報モザイク済の画面をここに貼付）", base_size=9)

def parse_table(lines, i, doc, do_math=False):
    rows = []
    while i < len(lines) and lines[i].lstrip().startswith('|'):
        rows.append(lines[i].strip()); i += 1
    cells = [[c.strip() for c in re.split(r'(?<!\\)\|', r)[1:-1]] for r in rows]
    cells = [c for c in cells if not (c and set(''.join(c).replace(' ','')) <= set('-:'))]  # 区切り行除去
    if not cells:
        return i
    ncol = max(len(r) for r in cells)
    t = doc.add_table(rows=len(cells), cols=ncol)
    t.style = 'Light Grid Accent 1'
    for ri, row in enumerate(cells):
        for ci in range(ncol):
            cell = t.rows[ri].cells[ci]
            cell.paragraphs[0].clear()
            add_inline(cell.paragraphs[0], row[ci] if ci < len(row) else "", do_math=do_math, base_size=9)
            if ri == 0:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
    doc.add_paragraph()
    return i

def convert(md_path, docx_path, do_math=False):
    with open(md_path, encoding='utf-8') as f:
        lines = f.read().split('\n')
    doc = Document()
    # 既定フォント(日本語)
    normal = doc.styles['Normal']; normal.font.name = JP_FONT; normal.font.size = Pt(10.5)
    normal.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), JP_FONT)
    for hn in ('Heading 1','Heading 2','Heading 3','Heading 4','Title'):
        try:
            st = doc.styles[hn]; st.font.name = JP_FONT
            st.element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), JP_FONT)
            st.font.color.rgb = RGBColor(0x1F,0x39,0x55)
        except KeyError:
            pass
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if s == '':
            i += 1; continue
        # コードフェンス
        if s.startswith('```'):
            i += 1; code = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code.append(lines[i]); i += 1
            i += 1
            p = doc.add_paragraph(); shade(p, "F2F3F5")
            for j, cl in enumerate(code):
                if j: p.add_run().add_break()
                r = p.add_run(cl); set_run_font(r, size=8.5, mono=True)
            continue
        # 表示数式ブロック \[ ... \]
        if s.startswith('\\['):
            inner=[]
            if '\\]' in s and s != '\\[':
                inner_text = s; i += 1
            else:
                i += 1
                while i < len(lines) and '\\]' not in lines[i]:
                    inner.append(lines[i]); i += 1
                if i < len(lines): i += 1
                inner_text = ' '.join(inner)
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
            shade(p, "F7F8FA")
            add_inline(p, inner_text, do_math=True, base_size=11)
            continue
        # 図プレースホルダ
        mfig = re.match(r'\[\[FIGURE:\s*([^\|]+)\|\s*(.+?)\]\]', s)
        if mfig:
            add_figure(doc, mfig.group(1).strip(), mfig.group(2).strip()); i += 1; continue
        # 水平線(細い罫線)
        if re.match(r'^---+$', s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bo = OxmlElement('w:bottom')
            bo.set(qn('w:val'),'single'); bo.set(qn('w:sz'),'6'); bo.set(qn('w:space'),'1'); bo.set(qn('w:color'),'C8CCD2')
            pbdr.append(bo); pPr.append(pbdr)
            i += 1; continue
        # 見出し
        mh = re.match(r'^(#{1,6})\s+(.*)$', s)
        if mh:
            lvl = len(mh.group(1)); h = doc.add_heading(level=min(lvl,4))
            add_inline(h, mh.group(2), do_math, base_size=16-lvl*1.5 if lvl<=2 else 12)
            i += 1; continue
        # 表
        if s.startswith('|'):
            i = parse_table(lines, i, doc, do_math=do_math); continue
        # 引用
        if s.startswith('>'):
            quote = []
            while i < len(lines) and lines[i].strip().startswith('>'):
                quote.append(re.sub(r'^\s*>\s?','',lines[i])); i += 1
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.3)
            shade(p, "EEF3F8")
            add_inline(p, ' '.join(q for q in quote if q.strip()), do_math, base_size=10)
            continue
        # 箇条書き
        mb = re.match(r'^(\s*)[-*]\s+(.*)$', line)
        if mb:
            p = doc.add_paragraph(style='List Bullet')
            add_inline(p, mb.group(2), do_math); i += 1; continue
        mn = re.match(r'^(\s*)\d+\.\s+(.*)$', line)
        if mn:
            p = doc.add_paragraph(style='List Number')
            add_inline(p, mn.group(2), do_math); i += 1; continue
        # 通常段落
        p = doc.add_paragraph(); add_inline(p, s, do_math); i += 1
    doc.save(docx_path)
    print("wrote", docx_path)

if __name__ == '__main__':
    # 紹介資料は PowerPoint(_build_pptx.py)が正。Word はロジック詳説(読む用)のみ生成する。
    jobs = [
        ("kiduri-補足_生成ロジックのアカデミック詳説.md", "kiduri-補足_生成ロジックのアカデミック詳説.docx", True),
    ]
    for md, dx, math in jobs:
        convert(os.path.join(HERE, md), os.path.join(HERE, dx), do_math=math)
